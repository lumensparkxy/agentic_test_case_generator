#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "client_submission"
SCREENSHOT_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "client_demo_agentic_testing_solution_brief.docx"

INK = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(91, 104, 121)


def set_run(run, size: float = 11, color: RGBColor = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run(run, 16 if level == 1 else 13, BLUE, True)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    set_run(paragraph.add_run(text), 10.6)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    set_run(paragraph.add_run(text), 10.4)


def output_location() -> str:
    try:
        return str(DOCX_PATH.relative_to(ROOT))
    except ValueError:
        return str(DOCX_PATH)


def add_metadata(doc: Document) -> None:
    rows = [
        ("Prepared for", "Synthetic client submission review"),
        ("Prepared by", "Demo Delivery Team"),
        ("Inputs", "Mocked browser workflow, synthetic requirements, optional generated screenshots"),
        ("Output location", output_location()),
    ]
    for label, value in rows:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run(paragraph.add_run(f"{label}: "), 10.3, INK, True)
        set_run(paragraph.add_run(value), 10.3)


def add_screenshot_gallery(doc: Document) -> None:
    screenshots = sorted(SCREENSHOT_DIR.glob("*.png"))[:8]
    if not screenshots:
        add_body(
            doc,
            "No screenshots were found. Run `node frontend/e2e/capture-client-screenshots.mjs` against a local frontend to populate the ignored client_submission/screenshots directory.",
        )
        return

    for item in screenshots:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(paragraph.add_run(item.name), 9.5, MUTED, True)
        try:
            paragraph.add_run().add_picture(str(item), width=Inches(6.25))
        except Exception as exc:
            add_body(doc, f"Skipped screenshot {item.name}: {exc}")


def build_docx() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(title.add_run("Agentic Testing Client Demo Brief"), 22, INK, True)

    subtitle = doc.add_paragraph()
    set_run(
        subtitle.add_run("Synthetic workflow package for reviewing requirements-to-test-case generation, grounded context, traceability, diagnostics, and export readiness."),
        12,
        MUTED,
        italic=True,
    )
    add_metadata(doc)

    add_heading(doc, "Scope", 1)
    add_body(
        doc,
        "This brief packages a reproducible client-facing demonstration without committing generated screenshots, exports, private browser profiles, credentials, or real operational data.",
    )
    add_bullet(doc, "The capture script uses mocked API responses and synthetic investment-management requirements.")
    add_bullet(doc, "Generated artifacts are written under client_submission/, which is ignored by git.")
    add_bullet(doc, "The brief builder can include screenshots when they exist, but it also runs without them.")

    add_heading(doc, "Demonstrated Workflow", 1)
    add_bullet(doc, "Sign-in surface and authenticated demo session.")
    add_bullet(doc, "Requirement parse and approval gate with quality diagnostics.")
    add_bullet(doc, "Grounded context analysis for app, API, and workflow facts.")
    add_bullet(doc, "Generated test cases with traceability, requirement analysis, coverage, and workflow diagnostics.")
    add_bullet(doc, "Export readiness without committing downloaded artifacts.")

    add_heading(doc, "Reproducibility", 1)
    add_body(doc, "Start the frontend locally, then run:")
    add_bullet(doc, "cd frontend && npm run dev")
    add_bullet(doc, "node frontend/e2e/capture-client-screenshots.mjs")
    add_bullet(doc, "python scripts/build_client_solution_brief.py")

    add_heading(doc, "Screenshot Gallery", 1)
    add_screenshot_gallery(doc)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
